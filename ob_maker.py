import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import argparse

# Function to set font for the entire table
def set_table_font(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                    run.font.size = Pt(10)

# Function to set column widths
def set_column_widths(table, col_widths):
    for i, width in enumerate(col_widths):
        for row in table.rows:
            cell = row.cells[i]
            cell.width = width

# Function to print a title or ASCII art
def print_title():
    title = """
    ________             __                 __
   /  _____/_____     __|  | ____    ____ |  | __ __  ____
  /   \  __\__  \   / __ |/ __ \  / ___\|  | |  |/ ___\
  \    \_\  \/ __ \_/ /_/ \  ___/ / /_/  >  |_|  \  \___
   \______  (____  /\____/ \___  >___  /  /____/  /\___  /
          \/    \/      \/     \/     \/          \/
    Script created by Sanskar Kalra
    """
    print(title)

# Print the title or ASCII art
print_title()

# Parse command-line arguments
parser = argparse.ArgumentParser(description='Convert Excel data to Word table')
parser.add_argument('-e', '--excel', required=True, help='Path to the Excel file')
parser.add_argument('-w', '--word', required=True, help='Path to save the output Word file')
args = parser.parse_args()

# Load the Excel file
excel_file = args.excel
df = pd.read_excel(excel_file)

# Print the column names to verify
print(df.columns)

# Define the mapping between Excel headers and Word headers
header_mapping = {
    'Sr No.': 'Sr No',
    'Vulnerability ID(CVE/CWE)': 'CVE/CWE',
    'Vulnerability Name': 'Observation/ Vulnerability Title',
    'Risk Severity ': 'Severity',  # Note the trailing space in 'Risk Severity '
    'Impact': 'Impact',
    'Remediation': 'Recommendation',
    'Asset Details': 'Affected Endpoint',
    'Reference': 'Reference'
}

# Define the extra rows to be added (excluding the ones already mapped)
extra_rows = [
    'Detailed observation / Vulnerable point',
    'New or Repeat observation',
    'Management Comment',
    'Final Status'
]

# Combined list to ensure specific order
combined_order = [
    'Sr No',
    'Affected Endpoint',
    'Observation/ Vulnerability Title',
    'Detailed observation / Vulnerable point',
    'CVE/CWE',
    'Impact',
    'Severity',
    'Recommendation',
    'Reference',
    'New or Repeat observation',
    'Management Comment',
    'Final Status'
]

# Create a new Word document
doc = Document()

# Set the page orientation to landscape
section = doc.sections[0]
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height

# Add a title to the document
doc.add_heading('Vulnerability Report', level=1)

# Add data to the document, each observation in a new table
for index, row in df.iterrows():
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    table.style = 'Table Grid'

    # Set initial column widths (adjust as needed)
    col_widths = [Inches(1.5), Inches(6)]  # [header_width, details_width]
    set_column_widths(table, col_widths)

    # Add the headers to the first row of the table
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Header'
    hdr_cells[1].text = 'Details'

    # Add the data rows to the table in the specified order
    for word_header in combined_order:
        if word_header == 'New or Repeat observation':
            value = 'New observation'
        elif word_header == 'Final Status':
            value = 'Remediated and Closed ()'
        elif word_header in header_mapping.values():
            excel_header = list(header_mapping.keys())[list(header_mapping.values()).index(word_header)]
            value = row[excel_header]
        else:
            value = ''
        hdr_cells = table.add_row().cells
        hdr_cells[0].text = word_header
        hdr_cells[1].text = str(value)

    # Set font for the table
    set_table_font(table)

    # Add a paragraph for spacing
    doc.add_paragraph('\n')

# Save the Word document
output_file = args.word
doc.save(output_file)

print(f"Word document created successfully: {output_file}")
